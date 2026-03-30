"""Build Resume.docx — clean single-column, ATS-friendly layout."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT_BODY = "Calibri"
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_DARK_GRAY = RGBColor(0x33, 0x33, 0x33)


def set_run_font(run, size, bold=False, italic=False, color=COLOR_BLACK, small_caps=False):
    run.font.name = FONT_BODY
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    if small_caps:
        run.font.small_caps = True


def add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="999999"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=4)
    run = p.add_run(text.upper())
    set_run_font(run, 12, bold=True, small_caps=True, color=COLOR_DARK_GRAY)
    add_bottom_border(p)
    return p


def add_job_header(doc, title, company, date_range, location):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=2)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    run_title = p.add_run(title)
    set_run_font(run_title, 11, bold=True)
    run_sep = p.add_run(" — ")
    set_run_font(run_sep, 11)
    run_company = p.add_run(company)
    set_run_font(run_company, 11, bold=True)
    run_loc = p.add_run(f", {location}")
    set_run_font(run_loc, 11, italic=True, color=COLOR_DARK_GRAY)
    run_tab = p.add_run("\t")
    set_run_font(run_tab, 11)
    run_date = p.add_run(date_range)
    set_run_font(run_date, 11, color=COLOR_DARK_GRAY)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_run_font(run, 10)
    set_paragraph_spacing(p, before=1, after=1)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.35)
    pf.first_line_indent = Inches(-0.2)
    return p


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(10)
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ═══════════════════════════════════════════
    # 1. HEADER
    # ═══════════════════════════════════════════
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_name, after=2)
    run = p_name.add_run("Flavio Smirne")
    set_run_font(run, 20, bold=True)

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_contact, after=4)
    run = p_contact.add_run("Orlando Metropolitan Area, FL  |  (XXX) XXX-XXXX  |  flavio@smirne.com  |  linkedin.com/in/smirne")
    set_run_font(run, 10, color=COLOR_DARK_GRAY)

    # ═══════════════════════════════════════════
    # 2. PROFESSIONAL SUMMARY
    # ═══════════════════════════════════════════
    add_section_heading(doc, "Professional Summary")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=2)
    summary = (
        "Data and engineering leader with 20+ years of experience building and scaling data organizations, defining enterprise data models, and driving data-informed decision-making across financial services, retail, and e-commerce. "
        "Currently leading a 40-person organization at Appriss Retail spanning data engineering, software engineering, QA, and product delivery \u2014 responsible for the enterprise data model and data standards underpinning 40% of all U.S. omnichannel transactions across 60+ of the top 100 U.S. retailers. "
        "Architected the Appriss Retail Data Model v3, the foundational data structure powering all products across POS and e-commerce channels, and the basis for AI/ML model generation in return authorization decisioning."
    )
    run = p.add_run(summary)
    set_run_font(run, 10)

    # ═══════════════════════════════════════════
    # 3. WORK EXPERIENCE
    # ═══════════════════════════════════════════
    add_section_heading(doc, "Work Experience")

    # ── Appriss Retail ──
    add_job_header(doc, "Senior Director of Implementation", "Appriss Retail", "12/2024 – Present", "Irvine, CA (Remote)")
    for b in [
        "Lead a 40-person organization across 8 direct reports (4 SWE Managers, 1 QA Manager, 1 Director of Data Engineering, 1 Product Owner, 1 Scrum Master), owning data engineering, software engineering, QA, and product delivery.",
        "Architected the Appriss Retail Data Model v3 \u2014 the baseline data structure for all products across POS and e-commerce channels, processing 40% of all U.S. omnichannel transactions and used by data science as the foundation for AI/ML model generation in return authorization decisioning.",
        "Own data standards for 60+ of the top 100 U.S. retailers, ensuring data reliability, governance, and consistency at enterprise scale.",
        "Led company-wide data standardization initiative, establishing consistent definitions across all business units and enabling self-service analytics and reporting for cross-functional stakeholders.",
        "Deployed AI-driven tooling for automated data discovery and QA validation, significantly reducing manual effort in ETL pipelines and data processing workflows.",
        "Reduced client go-live time by 40% through process redesign, workflow automation, and organizational streamlining.",
    ]:
        add_bullet(doc, b)

    # ── GameStop ──
    add_job_header(doc, "Director, E-Commerce Fraud", "GameStop", "03/2022 – 12/2024", "Grapevine, TX (Remote)")
    for b in [
        "Inherited an under-resourced fraud function and built it into a global fraud and analytics operation \u2014 hiring the team, selecting vendors, and standing up the technology stack from the ground up.",
        "Reduced fraud rates by 68% through ML-driven detection models, data-driven policy optimization, and real-time decisioning systems.",
        "Built a comprehensive fraud analytics and reporting infrastructure on GCP BigQuery, Amazon S3, and REST-based integrations \u2014 enabling data-driven decision-making across fraud, finance, and executive leadership.",
        "Partnered directly with the CFO, CISO, and CTO to embed data-informed risk management into organizational priorities and strategic planning.",
        "Built and automated the global chargeback mitigation and response programs, transforming a reactive process into a proactive, analytics-driven pipeline.",
        "Developed ML-based detection models and advanced analytics strategies targeting ATO, card testing, AML, and transaction fraud, materially reducing exposure across all e-commerce channels.",
    ]:
        add_bullet(doc, b)

    # ── Accertify ──
    add_job_header(doc, "Manager, Solution Productization", "Accertify (American Express)", "01/2014 – 03/2022", "Itasca, IL")
    for b in [
        "Defined industry-specific data models for retail, airline/travel, and ticketing verticals \u2014 foundational schemas deployed to every Accertify client globally and used as the basis for fraud detection, analytics, and reporting.",
        "Built a comprehensive analytics and reporting system that became a key differentiator in client retention and sales \u2014 enabling self-service insights for both internal teams and clients.",
        "Rewrote core detection algorithms, materially improving accuracy and reducing false positives through rigorous data analysis and experimentation.",
        "Designed reusable data-driven fraud detection and analytics components, establishing scalable patterns for data access, reporting, and decision support.",
        "Recognized with the President's Award for top 1% performers globally at American Express.",
        "Coordinated cross-functional teams across engineering, product, data science, and client services to deliver analytics and data platform tooling with zero downtime.",
    ]:
        add_bullet(doc, b)

    # ── State Farm ──
    add_job_header(doc, "Lead Developer, Mobile / Web", "State Farm Insurance", "04/2008 – 01/2014", "Bloomington, IL")
    for b in [
        "Architected the backend service aggregation portal that unified data feeds across all mobile platforms into a single API layer \u2014 a large-scale data integration and access project.",
        "Led the development of State Farm Pocket Agent for Windows Phone 7 and 8, one of the first mobile insurance apps in the industry.",
    ]:
        add_bullet(doc, b)

    # ── Realinked ──
    add_job_header(doc, "CTO", "Realinked.com", "02/2010 – 08/2011", "Chicago, IL")
    for b in [
        "Founded and built an online real estate brokerage from zero \u2014 designed the full data architecture, technology stack, and analytics infrastructure.",
        "Drove the platform to profitability and executed a successful exit within two years.",
    ]:
        add_bullet(doc, b)

    # ── Wolfram ──
    add_job_header(doc, "Software Engineering Manager", "Wolfram Research", "03/2004 – 04/2008", "Champaign, IL")
    for b in [
        "Promoted from tech support to software engineering manager within two years, recognized for technical depth and leadership potential.",
        "Designed and shipped the Image Processing, Import/Export, and 3D Graphics frameworks \u2014 core components still in Mathematica today.",
        "Led a team translating advanced mathematical research into production C++ across Windows, Linux, and macOS.",
    ]:
        add_bullet(doc, b)

    # ═══════════════════════════════════════════
    # 4. SKILLS
    # ═══════════════════════════════════════════
    add_section_heading(doc, "Skills")

    leadership_skills = [
        "Data Org Building & Talent Development",
        "Enterprise Data Governance & Standards",
        "Cross-Functional Program Leadership",
        "Executive Stakeholder Partnership",
        "Data-Driven Decision-Making Culture",
        "Agile Transformation / Scrum at Scale",
        "Vendor & Partner Strategy",
        "Change Management & Process Redesign",
    ]
    technical_skills = [
        "Enterprise Data Modeling & Architecture",
        "AI/ML Model Enablement (LLMs, Traditional ML)",
        "Self-Service Analytics & Reporting Platforms",
        "Data Standardization & Quality at Scale",
        "Python, C#/.NET, SQL, JavaScript",
        "PostgreSQL, MS SQL, GCP BigQuery",
        "ETL Pipeline Design & Automation",
        "Tableau / Power BI / Splunk",
    ]

    max_rows = max(len(leadership_skills), len(technical_skills))
    table = doc.add_table(rows=max_rows + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for ci, header in enumerate(["Leadership & Operations", "Technical & Data"]):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, 10, bold=True)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    for ri in range(max_rows):
        left = leadership_skills[ri] if ri < len(leadership_skills) else ""
        right = technical_skills[ri] if ri < len(technical_skills) else ""
        for ci, text in enumerate([left, right]):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_run_font(run, 10)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # ═══════════════════════════════════════════
    # 5. EDUCATION
    # ═══════════════════════════════════════════
    add_section_heading(doc, "Education")

    for degree, school, year, note in [
        ("Bachelor of Science, Computer Engineering", "University of Illinois, Champaign-Urbana", "2003", "Dean's List"),
        ("Bachelor of Science, Civil Engineering", "University of São Paulo, São Carlos, SP, Brazil", "1999", ""),
    ]:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=2, after=2)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
        run = p.add_run(degree)
        set_run_font(run, 10, bold=True)
        run = p.add_run(f" — {school}")
        set_run_font(run, 10)
        if note:
            run = p.add_run(f" ({note})")
            set_run_font(run, 10, italic=True, color=COLOR_DARK_GRAY)
        run = p.add_run(f"\t{year}")
        set_run_font(run, 10, color=COLOR_DARK_GRAY)

    # ═══════════════════════════════════════════
    # 6. LANGUAGES
    # ═══════════════════════════════════════════
    add_section_heading(doc, "Languages")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=4)
    run = p.add_run("English  |  Portuguese  |  Spanish")
    set_run_font(run, 10)

    # ── Metadata ──
    props = doc.core_properties
    props.author = "Flavio Smirne"
    props.comments = ""

    # ── Save ──
    out_path = "./Resume.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
