"""Build Resume_Architect.docx — Chief / Principal Architect (IC) targeting.

Reuses the styling/layout helpers from build_resume.py (same visual identity)
and swaps in content framed around technical architecture, system design, and
org-wide technical strategy rather than people management.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_resume_director import (
    FONT_BODY,
    COLOR_BLACK,
    COLOR_DARK_GRAY,
    set_run_font,
    set_paragraph_spacing,
    add_section_heading,
    add_job_header,
    add_bullet,
)


def add_skill_group(doc, label, items):
    """A single 'Label: item, item, item' line — readable and ATS-friendly."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=1)
    run = p.add_run(f"{label}: ")
    set_run_font(run, 10, bold=True)
    run = p.add_run(", ".join(items))
    set_run_font(run, 10)
    return p


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(10)
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ═══════════════════════════════════════════
    # 1. HEADER
    # ═══════════════════════════════════════════
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_name, after=2)
    run = p_name.add_run("Flavio Smirne")
    set_run_font(run, 20, bold=True)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_title, after=2)
    run = p_title.add_run("Principal Data Architect")
    set_run_font(run, 11, bold=True, small_caps=True, color=COLOR_DARK_GRAY)

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_contact, after=4)
    run = p_contact.add_run(
        "Orlando Metropolitan Area, FL  |  (XXX) XXX-XXXX  |  flavio@smirne.com  |  linkedin.com/in/smirne"
    )
    set_run_font(run, 10, color=COLOR_DARK_GRAY)

    # ═══════════════════════════════════════════
    # 2. PROFESSIONAL SUMMARY
    # ═══════════════════════════════════════════
    add_section_heading(doc, "Professional Summary")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=2)
    summary = (
        "Hands-on data and platform architect with 20+ years designing enterprise data platforms, integration frameworks, "
        "real-time decisioning systems, and large-scale analytics architectures across retail, financial services, "
        "e-commerce, and insurance. Sets data architecture direction across entire product portfolios — author of canonical "
        "enterprise data models, semantic and governance frameworks, and re-platforming strategies adopted by 60+ of the top 100 U.S. "
        "retailers and underpinning ~40% of U.S. omnichannel transactions. Pairs deep data modeling and system design with the "
        "cross-organizational influence to align engineering, data science, product, and executives on a single "
        "technical strategy — owning the design decisions and trade-offs end to end."
    )
    run = p.add_run(summary)
    set_run_font(run, 10)

    # ═══════════════════════════════════════════
    # 3. CORE TECHNICAL EXPERTISE (front-loaded for IC/architect roles)
    # ═══════════════════════════════════════════
    add_section_heading(doc, "Core Technical Expertise")

    for label, items in [
        ("Architecture", [
            "Enterprise Data Architecture", "Solution & Integration Architecture",
            "Event-Driven & API Architecture", "Data Modeling & Information Architecture",
            "Platform Modernization & Re-Platforming",
        ]),
        ("Data & Platforms", [
            "Snowflake", "dbt", "BigQuery", "Azure Data Platforms", "PostgreSQL", "SQL Server",
            "ETL/ELT Pipeline Design", "Data Quality Frameworks", "Self-Service Analytics",
        ]),
        ("AI/ML & Decisioning", [
            "AI/ML Enablement (LLMs, traditional ML)", "Real-Time Decisioning & Scoring",
            "Fraud Decisioning Platforms", "Risk Scoring", "AML/KYC", "Identity & Account Protection",
        ]),
        ("Languages & Tools", [
            "Python", "C#/.NET", "SQL", "JavaScript", "REST APIs", "Tableau / Power BI / Splunk",
        ]),
        ("Governance & Practices", [
            "Data Governance & Standards", "Architecture Review & Standards-Setting",
            "Cross-Functional Technical Leadership", "Vendor & Technology Evaluation",
        ]),
    ]:
        add_skill_group(doc, label, items)

    # ═══════════════════════════════════════════
    # 4. WORK EXPERIENCE
    # ═══════════════════════════════════════════
    add_section_heading(doc, "Work Experience")

    # ── Appriss Retail ──
    add_job_header(doc, "Senior Director of Implementation", "Appriss Retail", "12/2024 – Present", "Irvine, CA (Remote)")
    for b in [
        "Own the enterprise technical architecture and data strategy across the full product portfolio — POS and e-commerce data spanning over $1 trillion in annual transactions and ~250M unique customer identifiers, representing ~40% of U.S. omnichannel transactions across 60+ of the top 100 U.S. retailers.",
        "Architected the Appriss Retail Data Model v3, the canonical enterprise data model serving as the single foundation for analytics, fraud detection, and AI/ML return-authorization decisioning across every product.",
        "Personally wrote the first POS and e-commerce parsers on ARDM v3 as a working proof-of-concept, validated against a live client across two delivery phases (POS, then e-commerce) — establishing the reference implementation that anchors the new ingestion model.",
        "Established enterprise data standards, semantic models, and governance frameworks that drive consistency and interoperability across product domains and client integrations — authoring core framework code that standardizes and validates ARDM v3 across all clients.",
        "Leading the re-architecture of the retail ingestion platform from Azure/Python pipelines feeding a Greenplum MPP cluster to a Snowflake/dbt ELT ecosystem — adopting Snowflake's separation of storage and compute to eliminate resource contention between ingestion load and analytics and to simplify elastic scaling across all clients.",
        "Built an LLM-backed data discovery tool (Claude) that auto-generates schema analysis of raw client data and produces a working conversion prototype mapping it to ARDM v3 — cutting client implementation timelines by 40%.",
        "Serve as principal technical advisor to engineering, product, customer success, and executive leadership on data strategy and platform architecture.",
    ]:
        add_bullet(doc, b)

    # ── GameStop ──
    add_job_header(doc, "Director, E-Commerce Fraud", "GameStop", "03/2022 – 12/2024", "Grapevine, TX (Remote)")
    for b in [
        "Re-architected GameStop's global e-commerce fraud platform around an ML-first decisioning model, relegating the legacy rules engine to a thin policy layer and cutting the maintained rule set by two-thirds.",
        "Reduced the fraud rate by 68% by shifting the bulk of decisioning from hand-tuned rules to ML models, with real-time scoring and data-driven policy optimization.",
        "Designed real-time fraud decisioning systems integrating ML models, third-party data providers, behavioral analytics, and risk signals into a unified scoring pipeline.",
        "Architected the fraud analytics and reporting platform on GCP BigQuery, Amazon S3, and REST integrations, enabling data-driven decisioning across fraud, finance, and executive leadership.",
        "Personally wrote much of the team's .NET tooling — including a bulk fraud-resolution client integrated directly with Accertify and the GameStop Dispute Management System, which generated specialized financial and analyst-performance reports and produced every fraud and chargeback report delivered to the C-suite.",
        "Designed fraud-intelligence data models for account-takeover detection, card-testing prevention, AML monitoring, chargeback analytics, and transaction-fraud prevention.",
    ]:
        add_bullet(doc, b)

    # ── Accertify ──
    add_job_header(doc, "Solution Productization Architect", "Accertify (American Express)", "01/2014 – 03/2022", "Itasca, IL")
    for b in [
        "Designed the airline data model from the ground up — the foundational schema behind fraud systems serving 8 of the 10 largest global airlines (United, Delta, Southwest, Air France-KLM) — and defined the retail and ticketing vertical models deployed to every Accertify client globally as the basis for fraud detection, analytics, and reporting.",
        "Architected platforms that sustained peak loads of 50M+ transactions and over $6B in processed volume on a single Black Friday.",
        "Re-engineered core fraud detection algorithms, materially improving accuracy and reducing false positives.",
        "Architected self-service analytics platforms that became a key differentiator in client retention and new-business acquisition.",
        "Delivered highly available enterprise fraud platforms with zero downtime, coordinating across engineering, data science, product, and client services.",
        "Recipient of the American Express President's Award for top 1% global performance.",
    ]:
        add_bullet(doc, b)

    # ── State Farm ──
    add_job_header(doc, "Lead Developer, Data & Mobile Platforms", "State Farm Insurance", "04/2008 – 01/2014", "Bloomington, IL")
    for b in [
        "Designed the data aggregation system that served as the central orchestrator for all data requests across mobile and web applications — a unified API layer consolidating data services from disparate enterprise backends.",
        "Led development of State Farm Pocket Agent for Windows Phone, one of the industry's earliest mobile insurance applications.",
    ]:
        add_bullet(doc, b)

    # ── Realinked ──
    add_job_header(doc, "CTO", "Realinked.com", "02/2010 – 08/2011", "Chicago, IL")
    for b in [
        "Designed the company's end-to-end technology architecture from zero — data platform, analytics infrastructure, and integration framework — scaling it through profitability to a successful exit.",
    ]:
        add_bullet(doc, b)

    # ── Wolfram ──
    add_job_header(doc, "Software Engineering Manager", "Wolfram Research", "03/2004 – 04/2008", "Champaign, IL")
    for b in [
        "Managed the Import/Export team, owning every data format supported in Mathematica — dozens of file formats across data ingestion and export.",
        "Personally wrote many of the file importers and authored the Rasterize function that converts Mathematica expressions into raster images — still core to the product today.",
    ]:
        add_bullet(doc, b)

    # ═══════════════════════════════════════════
    # 5. EDUCATION
    # ═══════════════════════════════════════════
    add_section_heading(doc, "Education")

    for degree, school, year, note in [
        ("Bachelor of Science, Computer Engineering", "University of Illinois, Champaign-Urbana", "2003", "Dean's List"),
        ("Bachelor of Science, Civil Engineering", "University of São Paulo, São Carlos, SP, Brazil", "1999", ""),
    ]:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=2, after=2)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
        run = p.add_run(degree)
        set_run_font(run, 10, bold=True)
        run = p.add_run(f" — {school}")
        set_run_font(run, 10)
        if note:
            run = p.add_run(f" ({note})")
            set_run_font(run, 10, italic=True, color=COLOR_DARK_GRAY)
        run = p.add_run(f"\t{year}")
        set_run_font(run, 10, color=COLOR_DARK_GRAY)

    # ═══════════════════════════════════════════
    # 6. LANGUAGES
    # ═══════════════════════════════════════════
    add_section_heading(doc, "Languages")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=4)
    run = p.add_run("English  |  Portuguese  |  Spanish")
    set_run_font(run, 10)

    # ── Metadata ──
    props = doc.core_properties
    props.author = "Flavio Smirne"
    props.comments = ""

    # ── Save ──
    out_path = "./Resume.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
